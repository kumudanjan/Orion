"""
RAG Engine — Retrieval Augmented Generation over Logs + SOPs
"""

import logging
from typing import List, Dict, Any, Optional

from app.ingestion.vector_store import get_vector_store
from app.llm.client import get_llm_client
from app.guardrails.pii_masker import validate_llm_output, hallucination_check

logger = logging.getLogger(__name__)

_RAG_SYSTEM = """You are a helpful SRE assistant. You are given retrieved log snippets and/or SOP documentation as context.
Answer the user's question using ONLY the provided context. If the answer is not in the context, say so clearly.
Be concise and precise. Never reveal credentials, tokens, or sensitive data."""


def rag_query(
    question: str,
    collections: List[str] = None,
    top_k: int = 5,
) -> Dict[str, Any]:
    """
    Run RAG pipeline:
      1. Retrieve top-k chunks from vector store
      2. Build prompt with context
      3. Call LLM
      4. Validate output (RAI check + hallucination score)
    """
    if collections is None:
        collections = ["logs", "sops"]

    store = get_vector_store()
    client = get_llm_client()

    # Retrieve from all collections
    all_chunks: List[Dict[str, Any]] = []
    for col in collections:
        hits = store.query(question, collection=col, top_k=top_k)
        for h in hits:
            h["collection"] = col
        all_chunks.extend(hits)

    # Sort by score, keep top_k overall
    all_chunks.sort(key=lambda x: x.get("score", 0), reverse=True)
    all_chunks = all_chunks[:top_k]

    context_text = "\n\n---\n\n".join(
        f"[Source: {c.get('collection','?')} | Score: {c.get('score',0):.2f}]\n{c['text']}"
        for c in all_chunks
    )

    if client is None:
        return {
            "answer": "LLM not configured. Please add your API key to .env.",
            "sources": all_chunks,
            "is_safe": True,
            "hallucination_score": 0.0,
        }

    prompt = f"Context:\n{context_text}\n\nQuestion: {question}"
    answer = client.complete(system=_RAG_SYSTEM, user=prompt)

    # RAI checks
    is_safe, violations = validate_llm_output(answer)
    if not is_safe:
        logger.warning("RAG answer failed safety check — redacting")
        answer = "[Response blocked by RAI guardrail due to potential sensitive data exposure.]"

    h_score = hallucination_check(answer, [c["text"] for c in all_chunks])

    return {
        "answer": answer,
        "sources": all_chunks,
        "is_safe": is_safe,
        "violations": violations,
        "hallucination_score": h_score,
    }

from docx import Document
# import docx
import io

def extract_text_from_docx(file_like) -> str:
    """
    file_like: a file path, bytes, or a file-like object (e.g., Streamlit's UploadedFile)
    Returns: full text with paragraphs and table rows separated by newlines.
    """
    # If file_like is a file-like stream (e.g., UploadedFile), wrap in BytesIO
    doc = Document(io.BytesIO(file_like.read()) if hasattr(file_like, "read") else file_like)

    parts = []

    # Paragraphs
    for p in doc.paragraphs:
        parts.append(p.text)

    # Tables (optional)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))

    text = "\n".join(parts).strip()
    return text

from pypdf import PdfReader

def extract_text_from_pdf_pypdf2(file_like) -> str:
    reader = PdfReader(file_like)  # accepts a path or file-like object
    texts = []
    for page in reader.pages:
        texts.append(page.extract_text() or "")
    return "\n\n".join(texts).strip()

def ingest_sop_document(text: str, title: str = "SOP") -> int:
    """
    Ingest a standard operating procedure document into the 'sops' vector collection.
    Splits by paragraph, embeds, and upserts.
    """
    store = get_vector_store()
    chunk_size = 1000
    overlap = 100
    max_chunks = 100
    text_chunks = []
    start = 0
    while (start < len(text) ):
        end = start + chunk_size
        chunk = text[start:end].strip()
        if chunk:
            text_chunks.append(chunk)
            start+=chunk_size - overlap
        if len(text_chunks)>=max_chunks:
            break
    chunks = [{"text":c,"metadata":{"title":title,"type":"sop"}}for c in text_chunks]
    batch_size = 25
    total = 0
    for i in range(0,len(chunks),batch_size):
        batch = chunks[i:i+batch_size]
        total += store.upsert(batch,collection="sops")
    logger.info("Ingested %d SOP chunks for '%s'", total, title)
    return total
